# coding:utf-8
# author:Super
from docx import Document
import random
import os

# def cal(n, st=1):
#     set1 = set()
#     while True:
#         a = random.randint(st, 9)
#         b = random.randint(st, 9)
#
#         str = "{} x {}".format(a, b)
#
#         if str[::-1] not in set1:
#             set1.add(str)
#
#         if len(set1) > n - 1:
#             break
#
#     count = 1
#     str = ""
#     for i in set1:
#         if count % 4 == 0:
#             str += "{} =    \n".format(i)
#         else:
#             str += "{} =     ".format(i)
#         count += 1
#
#     return str

def cal(n, st=1):
    set1 = set()
    while True:
        a = random.randint(st, 9)
        b = random.randint(st, 9)

        str = "{}  ×  {}".format(a, b)

        if str[::-1] not in set1:
            set1.add(str)

        if len(set1) > n - 1:
            break

    count = 1
    str = ""
    for i in set1:
        if count % 4 == 0:
            str += "{}  =                   \n\n".format(i)
        else:
            str += "{}  =                   ".format(i)
        count += 1

    return str


# # ================输出一个文件=====================
# # 创建word文档对象
# document = Document()
#
# # 添加标题
# document.add_heading('数学试卷', 0)
#
# # 添加数据
# p = document.add_paragraph(cal(20, 2))
#
# # 保存world文档
# document.save(r'/Users/super/Desktop/demo.doc')

# print(os.listdir("/Users/super/Desktop/"))
# exit()

# 产出一个试卷的文件夹
if "试卷" not in os.listdir("/Users/super/Desktop/"):
    os.mkdir("/Users/super/Desktop/试卷")

for i in range(1, 101):
    document = Document()

    document.add_heading('数学试卷{}'.format(i), 0)

    p = document.add_paragraph(cal(20, 2))

    document.save(r'/Users/super/Desktop/试卷/试卷{}.doc'.format(i))

    print("第{}份数据处理完毕……".format(i))
